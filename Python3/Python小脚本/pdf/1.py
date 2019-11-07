# coding:utf-8
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator,TextConverter
from pdfminer.layout import LAParams,LTTextBoxHorizontal
from pdfminer.pdfpage import PDFTextExtractionNotAllowed,PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from docx import Document
import os,warnings
warnings.filterwarnings("ignore")

def ParsePDF():
    filename = open(pdfpath, 'rb')  #以二进制读模式打开
    #用文件对象来创建一个pdf文档分析器
    parser = PDFParser(filename)
    # 创建一个PDF文档对象存储文档结构,提供密码初始化，没有就不用传该参数
    doc = PDFDocument(parser, password='')
    #检查文件是否允许文本提取
    if not doc.is_extractable:
        print("Not Allowd Extractable")
        raise PDFTextExtractionNotAllowed
    
    # 创建PDf 资源管理器来管理共享资源，#caching = False不缓存
    rsrcmgr = PDFResourceManager(caching = False)
    # 创建一个PDF设备对象
    laparams = LAParams()
    # 创建一个PDF页面聚合对象
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    #device = TextConverter(rsrcmgr, retstr, codec='utf-8', laparams=laparams)
    # 创建一个PDF解析器对象
    interpreter = PDFPageInterpreter(rsrcmgr, device)  

    # 获取page列表list对象,
    # print(PDFPage.get_pages(doc))

    #获取page列表循环遍历列表，每次处理一个page的内容
    for page in PDFPage.create_pages(doc):
        # 接受该页面的LTPage对象
        interpreter.process_page(page)
        # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
        layout = device.get_result()
        for i in layout:
            if hasattr(i,"get_text") :
                content = i.get_text().replace(u'\xa0',u'').replace('\n','')
                document.add_paragraph(content , style=None)
        break    
    document.save("a.docx")
    filename.close()
    return 1

if __name__ == "__main__":
    pdfpath = input("请输入pdf文件路径：")
    document = Document()
    ParsePDF()